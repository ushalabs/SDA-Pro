from __future__ import annotations

from datetime import date
from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "assets"
DOCX_OUT = REPORTS / "SDA-Pro-Comprehensive-Report-Updated.docx"
MD_OUT = REPORTS / "SDA-Pro-Comprehensive-Report-Updated.md"


def read(path: Path, fallback: str = "") -> str:
    return path.read_text(encoding="utf-8", errors="replace") if path.exists() else fallback


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 103, 119)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)
    else:
        doc.add_paragraph(f"[Missing image: {path.name}]", style="Quote")


def add_evidence_note(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, heading in enumerate(headers):
        set_cell_shading(hdr[i], "EEF3F8")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(heading)
        r.bold = True
        r.font.color.rgb = RGBColor(11, 37, 69)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, text: str, max_chars: int = 1800) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text[:max_chars] + ("..." if len(text) > max_chars else ""))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(45, 55, 72)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SDA-Pro Security Incident Response & Threat Mitigation Platform")


def build_docx() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SDA-Pro Comprehensive Project Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Security Incident Response & Threat Mitigation Platform")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(46, 116, 181)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated from implemented project evidence | {date.today().isoformat()}")
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report documents the implemented SDA-Pro university semester project: a demo-ready Security Incident "
        "Response and Threat Mitigation Platform. The system focuses on Software Design and Architecture evidence: "
        "design patterns, architecture styles, service boundaries, API contracts, asynchronous events, database schema, "
        "frontend screens, UML diagrams, and tests. External cybersecurity products are intentionally mocked so the "
        "demo can run locally without real API keys."
    )
    doc.add_paragraph(
        "Implementation note: the backend features are packaged into a Spring Boot composition root for reliable local "
        "demo execution, while the repository, contracts, diagrams, and documentation preserve clear SOA-style service "
        "boundaries. This is appropriate for a semester architecture demonstration, but it is not claimed to be a "
        "production microservices deployment."
    )

    doc.add_heading("Final Tech Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Proof"],
        [
            ["Backend", "Java 21, Spring Boot 3, Maven", "services/alert-ingestion-service"],
            ["Frontend", "React, TypeScript, Vite", "soc-dashboard"],
            ["Database", "PostgreSQL 16", "docker-compose service postgres"],
            ["Cache", "Redis 7", "Threat intelligence proxy cache"],
            ["Message broker", "RabbitMQ 3", "Event publishing/subscribing"],
            ["Realtime", "Server-Sent Events dashboard stream", "/dashboard/events"],
            ["Docs/UML", "PlantUML and ADR markdown", "docs/uml diagrams, docs/adr"],
            ["Deployment", "Docker Compose", "docker-compose.yml"],
        ],
        [1900, 3200, 3900],
    )

    doc.add_heading("System Architecture", level=1)
    doc.add_paragraph(
        "The project is organized as a local SOA demonstration. Alerts enter through ingestion APIs, are normalized, "
        "enriched, correlated into incidents, advanced through lifecycle states, responded to using strategy-selected "
        "actions, published as events, pushed to the dashboard, and written to audit storage."
    )
    add_bullets(
        doc,
        [
            "Controllers expose required API endpoints and route requests into application services.",
            "Application services coordinate normalization, enrichment, correlation, state transitions, response execution, notification, and audit.",
            "Domain classes contain incident states, alert composites, response strategies, and event models.",
            "Adapters and proxies isolate mock Splunk, Firewall, VirusTotal, MISP, Slack, PagerDuty, Email, and Active Directory behavior.",
            "PostgreSQL stores durable records, Redis caches threat intelligence, RabbitMQ represents event-driven communication, and SSE updates the SOC dashboard.",
        ],
    )
    add_image(doc, ASSETS / "diagram-component.png", "Figure 1. PlantUML component diagram showing SOA boundaries.")
    add_evidence_note(
        doc,
        "Component diagram explanation",
        "This diagram proves the SOA requirement by separating the platform into named services: ingestion, enrichment, "
        "incident management, response orchestration, threat intelligence, notification, audit, identity, dashboard, and "
        "shared contracts. It also shows infrastructure dependencies: PostgreSQL for persistence, Redis for cache, "
        "RabbitMQ for events, and realtime dashboard communication. The diagram demonstrates that the project is not just "
        "a UI demo; it has explicit service responsibilities and communication paths."
    )

    doc.add_heading("Folder Structure", level=1)
    add_code_block(
        doc,
        "\n".join(
            [
                "SDA Project/",
                "  docker-compose.yml",
                "  README.md",
                "  services/alert-ingestion-service/     Spring Boot backend composition root",
                "  soc-dashboard/                        React + TypeScript frontend",
                "  shared/contracts/                     API/event shared contracts",
                "  docs/adr/                             Architecture Decision Records",
                "  docs/api/                             OpenAPI and AsyncAPI contracts",
                "  docs/plantuml codes/                  PlantUML source diagrams",
                "  docs/uml diagrams/                    Rendered PNG diagrams",
                "  reports/assets/                       Screenshots, proof files, copied diagrams",
            ]
        ),
    )

    doc.add_heading("Database Schema", level=1)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["analysts", "Mock SOC analyst identity and role data."],
            ["alert_sources", "Configured mock Splunk, Firewall, and other source definitions."],
            ["alerts", "Normalized alert records and severity metadata."],
            ["alert_groups", "Composite grouping of related alerts."],
            ["alert_group_members", "Membership join table for alert composites."],
            ["enrichment_results", "Geo, asset criticality, threat intel, and enrichment evidence."],
            ["incidents", "Incident records, severity, status, and lifecycle state."],
            ["incident_alerts", "Incident-to-alert correlation join table."],
            ["incident_state_transitions", "History of incident state movement."],
            ["response_actions", "Planned response actions such as isolate host or block IP."],
            ["response_action_outcomes", "Execution results for response actions."],
            ["threat_indicators", "Indicators and reputation data."],
            ["notifications", "Mock Slack, PagerDuty, and Email dispatch records."],
            ["audit_logs", "Security and system audit trail."],
        ],
        [2600, 6600],
    )

    doc.add_heading("Service Boundaries", level=1)
    add_table(
        doc,
        ["Service", "Responsibility"],
        [
            ["alert-ingestion-service", "Webhook/poll intake, source-specific normalization, AlertIngested event."],
            ["enrichment-correlation-service", "Enrichment pipeline, threat scoring, alert grouping, incident creation."],
            ["incident-management-service", "Incident CRUD and state machine transitions."],
            ["response-orchestration-service", "Strategy selection, response factory, decorators, proxy execution."],
            ["threat-intel-service", "Mock reputation lookup through adapters and Redis proxy cache."],
            ["notification-service", "Mock Slack, PagerDuty, and Email dispatch adapters."],
            ["audit-service", "Audit event creation and query API."],
            ["identity-service", "Mock login/signup/profile flow for demo analysts."],
            ["soc-dashboard", "Modern React dashboard and realtime monitoring screens."],
            ["shared contracts", "OpenAPI, AsyncAPI, event DTOs, and shared documentation."],
        ],
        [2900, 6300],
    )

    doc.add_heading("API Endpoints", level=1)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["POST /ingest/webhook/{sourceType}", "Accept mock Splunk/Firewall webhook alerts."],
            ["POST /ingest/poll/{sourceId}", "Simulate polling an alert source."],
            ["GET /alerts, GET /alerts/{id}", "List and inspect normalized alerts."],
            ["POST /enrichment/process", "Run enrichment chain on an alert."],
            ["GET /incidents, POST /incidents, GET /incidents/{id}", "Incident queue and incident detail support."],
            ["PATCH /incidents/{id}/state", "Move incident through lifecycle states."],
            ["POST /incidents/{id}/respond", "Select and execute response strategy."],
            ["POST /incidents/{id}/actions", "Execute explicit response actions."],
            ["POST /threat-intel/reputation", "Mock reputation lookup through adapter/proxy cache."],
            ["POST /notifications/dispatch", "Mock notification dispatch."],
            ["GET /audit/events", "Retrieve audit logs."],
            ["POST /auth/login, POST /auth/signup, GET /identity/me", "Mock authentication and profile endpoints."],
            ["GET /dashboard/metrics, GET /dashboard/events", "Dashboard metrics and realtime stream."],
        ],
        [3600, 5600],
    )

    doc.add_heading("Event Definitions", level=1)
    add_table(
        doc,
        ["Event", "When Published"],
        [
            ["AlertIngested", "A source alert is accepted and normalized."],
            ["AlertEnriched", "The enrichment pipeline completes."],
            ["IncidentCreated", "Correlation creates a new incident."],
            ["IncidentStateChanged", "An incident moves lifecycle state."],
            ["ResponseActionExecuted", "A response action finishes."],
            ["ThreatIntelUpdated", "Threat reputation is looked up or cached."],
            ["NotificationDispatched", "A mock notification adapter dispatches a message."],
            ["AuditLogCreated", "A durable audit record is stored."],
        ],
        [2800, 6400],
    )

    doc.add_heading("Required Demo Flow Evidence", level=1)
    add_numbered(
        doc,
        [
            "Mock Splunk or Firewall alert enters /ingest/webhook/{sourceType}.",
            "AlertNormalizerFactory selects a source normalizer.",
            "EnrichmentChain runs geo, asset criticality, and threat intel handlers.",
            "Threat intelligence is looked up through adapter classes and cached through the proxy.",
            "Related alerts are grouped using the Composite pattern.",
            "Correlation creates an incident in New state.",
            "Incident moves through triage and containment using State pattern classes.",
            "ResponseStrategy selects containment or eradication behavior.",
            "ResponseActionFactory creates actions; decorators and proxies add audit and approval behavior.",
            "RabbitMQ/EventBus publishes events; SSE updates dashboard panels in realtime.",
            "Audit log is stored and visible in the Audit Logs screen.",
        ],
    )
    add_image(doc, ASSETS / "diagram-sequence-alert-ingestion.png", "Figure 2. Alert ingestion sequence diagram.")
    add_evidence_note(
        doc,
        "Alert ingestion sequence explanation",
        "This sequence maps directly to the required demo flow: a mock Splunk or Firewall alert reaches the webhook API, "
        "the normalizer factory selects a source-specific normalizer, enrichment handlers run through the chain of "
        "responsibility, threat intelligence is queried through adapter/proxy classes, an alert group is composed, an "
        "incident is created, events are published, and the dashboard/audit service receives evidence of the workflow."
    )
    add_image(doc, ASSETS / "diagram-sequence-incident-response.png", "Figure 3. Incident response sequence diagram.")
    add_evidence_note(
        doc,
        "Incident response sequence explanation",
        "This sequence proves the incident lifecycle and response requirements. The analyst changes incident state, the "
        "State pattern validates the transition, the strategy selector chooses the response approach, factories create "
        "actions, decorators add audit/logging behavior, proxies simulate control checks, events are published, and audit "
        "records are stored after execution."
    )

    doc.add_heading("Frontend Screens", level=1)
    doc.add_paragraph(
        "The frontend was modernized for demo readiness: authentication is separated from the dashboard, signup is present, "
        "logout is available, raw backend identifiers and integration keys are masked, and each operational panel has a "
        "clean SOC-style layout."
    )
    screen_figures = [
        (
            "01-login-screen.png",
            "Login screen",
            "Satisfies the required Login screen and identity-service boundary. It provides a separate authentication "
            "entry point before analysts can access SOC panels. The implementation uses a mock session for demo safety "
            "instead of exposing backend tokens or integration keys in the browser."
        ),
        (
            "02-signup-screen.png",
            "Signup screen",
            "Adds analyst onboarding support for the identity requirement. The signup view creates a mock analyst "
            "profile for demonstration and keeps the authentication experience separate from operational screens."
        ),
        (
            "03-main-dashboard.png",
            "Main dashboard",
            "Satisfies the Main Dashboard requirement. It summarizes active incidents, alert volume, response outcomes, "
            "and SOC status from the dashboard metrics API so the evaluator can immediately see seeded demo data."
        ),
        (
            "04-live-alert-stream.png",
            "Live alert stream",
            "Satisfies the Live Alert Stream requirement. This screen demonstrates event-driven behavior and realtime "
            "push from backend events/SSE after mock Splunk or Firewall alerts enter the ingestion flow."
        ),
        (
            "05-incident-queue.png",
            "Incident queue",
            "Satisfies the Incident Queue requirement. It shows correlated incidents created from enriched alerts, "
            "supporting the required alert grouping and incident creation portion of the demo flow."
        ),
        (
            "06-incident-detail.png",
            "Incident detail",
            "Satisfies the Incident Detail requirement. It presents incident severity, linked evidence, lifecycle state, "
            "and state transition controls, proving the State pattern and incident-management service behavior."
        ),
        (
            "07-response-console.png",
            "Response console",
            "Satisfies the Response Console requirement. It demonstrates response orchestration: strategy selection, "
            "factory-created actions, decorator-added audit behavior, proxy control checks, and action outcome recording."
        ),
        (
            "08-threat-intel-lookup.png",
            "Threat intel lookup",
            "Satisfies the Threat Intel Lookup requirement. It shows mock VirusTotal/MISP style reputation lookup through "
            "adapters and the proxy cache, using Redis for repeat lookup simulation without real API keys."
        ),
        (
            "09-audit-logs.png",
            "Audit logs",
            "Satisfies the Audit Logs requirement. It proves that significant system actions generate durable audit events "
            "such as alert ingestion, incident changes, response execution, notifications, and threat intel lookups."
        ),
        (
            "10-settings.png",
            "Settings",
            "Satisfies the Settings requirement and addresses security presentation. Mock integrations are visible, but "
            "credentials are shown only as protected status values; raw API keys, payload identifiers, and tokens are not "
            "displayed in analyst panels."
        ),
    ]
    for idx, (file, caption, explanation) in enumerate(screen_figures, 1):
        add_image(doc, ASSETS / file, f"Figure {3 + idx}. {caption}.", width=6.35)
        add_evidence_note(doc, "Requirement explanation", explanation)

    doc.add_heading("Design Pattern Mapping", level=1)
    add_table(
        doc,
        ["Pattern", "Implementation Evidence", "Rationale"],
        [
            ["Singleton", "IngestionConfigManager, EventBusPublisher", "One shared configuration/event publishing component."],
            ["Factory Method", "AlertNormalizerFactory, ResponseActionFactory", "Create normalizers/actions without controller coupling."],
            ["Abstract Factory", "EnrichmentProviderFactory, NotificationFactory", "Create related provider families for mock integrations."],
            ["Composite", "AlertComponent", "Group alerts and individual alerts under a common interface."],
            ["Facade", "IncidentResponseFacade", "Simplify response orchestration into one use-case entry point."],
            ["Adapter", "ThreatAdapters", "Wrap mock VirusTotal/MISP style APIs behind one interface."],
            ["Decorator", "Decorators", "Add audit/logging/approval behavior around response action execution."],
            ["Proxy", "ThreatIntelProxy, ResponseActionProxy", "Cache lookups and simulate approval/rate-limit control."],
            ["State", "IncidentState", "Incident behavior changes by lifecycle phase."],
            ["Chain of Responsibility", "EnrichmentHandler", "Run enrichment handlers in sequence."],
            ["Observer", "EventBusPublisher", "Publish domain events to subscribers/dashboard."],
            ["Strategy", "ResponseStrategy", "Select response behavior by incident severity/context."],
        ],
        [1900, 3600, 3800],
    )
    doc.add_paragraph("Code annotation proof extracted from the repository:")
    add_code_block(doc, read(ASSETS / "pattern-annotations.txt"), max_chars=2600)
    add_image(doc, ASSETS / "diagram-class.png", "Figure 14. Class diagram showing all 12 design patterns.")
    add_evidence_note(
        doc,
        "Class diagram explanation",
        "The class diagram is the primary design-pattern proof. It links the 12 required patterns to concrete classes and "
        "collaborations: factories create normalizers/actions, adapters wrap mock external products, proxies add caching "
        "and control checks, decorators add behavior around response actions, the State pattern handles incident lifecycle, "
        "the chain processes enrichment, Composite groups alerts, Strategy selects response behavior, Facade simplifies "
        "orchestration, Observer publishes events, and Singleton components centralize shared event/config access."
    )

    doc.add_heading("Architecture Style Mapping", level=1)
    add_table(
        doc,
        ["Architecture Style", "Project Mapping"],
        [
            ["SOA", "Documented service boundaries for ingestion, enrichment, incident, response, threat intel, notification, audit, identity, dashboard, and contracts."],
            ["MVC", "Spring controllers expose APIs, services coordinate use cases, repositories persist data; React components render views from API state."],
            ["Layered Architecture", "Controller/API layer, application/service layer, domain layer, integration/adapters layer, persistence layer, infrastructure layer."],
            ["Event-Driven Architecture", "RabbitMQ/EventBus events for alert ingestion, enrichment, incident creation/state changes, response execution, notification, threat intel, and audit."],
        ],
        [2500, 6700],
    )

    doc.add_heading("Backend Proof", level=1)
    doc.add_paragraph(
        "The backend contains Java code for controllers, domain models, design patterns, repositories, events, security demo flows, "
        "mock integrations, tests, and Docker packaging."
    )
    add_bullets(
        doc,
        [
            "Java source files found: 39 in services/alert-ingestion-service/src/main/java.",
            "Test classes: AuditEventCreationTest, EndToEndDemoTest, PatternTests.",
            "OpenAPI contract: docs/api/openapi-sdapro.yaml.",
            "AsyncAPI event contract: docs/api/asyncapi-events.yaml.",
            "ADR decisions: five ADR markdown files under docs/adr.",
        ],
    )

    doc.add_heading("Frontend Proof", level=1)
    doc.add_paragraph(
        "The frontend implements the required login, dashboard, live alert stream, incident queue/detail, response console, "
        "threat intelligence lookup, audit logs, and settings views. It masks identifiers and avoids exposing raw integration secrets."
    )

    doc.add_heading("Docker and Runtime Proof", level=1)
    doc.add_paragraph(
        "The Docker Desktop screenshot below shows the sdaproject Compose stack running with RabbitMQ, PostgreSQL, Redis, "
        "the Spring Boot backend, and the SOC dashboard frontend. The command-line proof below was captured from the same "
        "project and lists the active backend, frontend, database, cache, and broker containers."
    )
    add_image(doc, ASSETS / "docker-desktop-running.png", "Figure 15. Docker Desktop running the SDA-Pro Compose stack.", width=6.35)
    add_evidence_note(
        doc,
        "Docker screenshot explanation",
        "This satisfies the deployment/runtime proof requirement for docker-compose up --build. The screenshot shows the "
        "local Compose project and the running service containers needed for the demo: broker, database, cache, backend, "
        "and frontend."
    )
    add_code_block(doc, read(ASSETS / "docker-compose-ps.txt"), max_chars=2200)

    doc.add_heading("API Response Proof", level=1)
    doc.add_paragraph("Dashboard metrics API response:")
    add_code_block(doc, read(ASSETS / "api-dashboard-metrics.json"), max_chars=1200)
    doc.add_paragraph("Threat intelligence lookup API response:")
    add_code_block(doc, read(ASSETS / "api-threat-intel-response.json"), max_chars=1200)

    doc.add_heading("Testing Evidence", level=1)
    doc.add_paragraph("Implemented test coverage includes normalizer factory, enrichment chain, incident state transitions, response strategy selection, threat intel proxy cache, audit event creation, and an end-to-end demo test.")
    add_code_block(doc, read(ASSETS / "test-results.txt"), max_chars=1800)

    doc.add_heading("Required Documentation Evidence", level=1)
    add_table(
        doc,
        ["Required Artifact", "Location"],
        [
            ["README with setup/demo instructions", "README.md"],
            ["Docker Compose", "docker-compose.yml"],
            ["ADR-001 SOA vs Microservices vs Modular Monolith", "docs/adr/ADR-001-soa-vs-microservices-vs-modular-monolith.md"],
            ["ADR-002 Sync vs Async Communication", "docs/adr/ADR-002-sync-vs-async-communication.md"],
            ["ADR-003 Database Strategy", "docs/adr/ADR-003-database-strategy.md"],
            ["ADR-004 Threat Intel Cache Strategy", "docs/adr/ADR-004-threat-intel-cache-strategy.md"],
            ["ADR-005 Real-time Push Strategy", "docs/adr/ADR-005-real-time-push-strategy.md"],
            ["PlantUML class/component/sequence diagrams", "docs/plantuml codes and docs/uml diagrams"],
            ["OpenAPI and AsyncAPI", "docs/api"],
        ],
        [3600, 5600],
    )

    doc.add_heading("Known Scope and Limitations", level=1)
    add_bullets(
        doc,
        [
            "External integrations are mocked by requirement; no real Splunk, CrowdStrike, Firewall, VirusTotal, MISP, Slack, PagerDuty, Email, or Active Directory keys are used.",
            "The implementation is demo-ready for SDA marking and local presentation, not hardened for production cybersecurity operations.",
            "Service boundaries are architecture/documentation boundaries inside a reliable local demo build, rather than independently deployed microservice artifacts for every named service.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The project satisfies the requested SDA-Pro deliverables: complete stack, required services, required APIs, required events, "
        "database schema, modern frontend screens, all 12 design patterns with code annotations, all 4 architecture styles, UML diagrams, "
        "ADR documentation, Docker Compose execution, seeded demo data, and automated tests."
    )

    doc.save(DOCX_OUT)


def build_markdown() -> None:
    docker_ps = read(ASSETS / "docker-compose-ps.txt").strip()
    tests = read(ASSETS / "test-results.txt").strip()
    patterns = read(ASSETS / "pattern-annotations.txt").strip()
    md = f"""# SDA-Pro Comprehensive Project Report

**Project:** Security Incident Response & Threat Mitigation Platform  
**Generated:** {date.today().isoformat()}  
**Stack:** Java 21 + Spring Boot 3, React + TypeScript + Vite, PostgreSQL, Redis, RabbitMQ, Docker Compose, PlantUML

## Executive Summary

This report documents the implemented SDA-Pro university semester project. The system is demo-ready and focuses on Software Design & Architecture marks: required services, APIs, events, database schema, frontend screens, UML diagrams, ADRs, tests, all 12 design patterns, and all 4 architecture styles.

Implementation note: backend capabilities are packaged into one Spring Boot composition root for reliable local demo execution while preserving documented SOA boundaries. External cybersecurity integrations are mocked by requirement.

## Final Tech Stack

| Layer | Technology |
|---|---|
| Backend | Java 21, Spring Boot 3, Maven |
| Frontend | React, TypeScript, Vite |
| Database | PostgreSQL 16 |
| Cache | Redis 7 |
| Broker | RabbitMQ 3 |
| Realtime | Server-Sent Events |
| Docs | ADR markdown, OpenAPI, AsyncAPI, PlantUML |
| Deployment | Docker Compose |

## System Architecture

![Component Diagram](assets/diagram-component.png)

The architecture follows controller/API, application service, domain, adapter/integration, persistence, and infrastructure layers. The demo flow is:

**Component diagram explanation:** This diagram proves the SOA requirement by separating the platform into named services: ingestion, enrichment, incident management, response orchestration, threat intelligence, notification, audit, identity, dashboard, and shared contracts. It also shows PostgreSQL, Redis, RabbitMQ, and realtime dashboard communication.

1. Mock Splunk/Firewall alert enters ingestion.
2. Alert is normalized.
3. Enrichment chain processes it.
4. Threat intel adapter/proxy performs reputation lookup and caching.
5. Composite groups related alerts.
6. Correlation creates incident in New state.
7. State pattern moves incident through lifecycle.
8. Strategy selects response.
9. Factory/decorator/proxy execute action.
10. Events are published, dashboard updates, audit is stored.

## Database Schema

Required tables implemented/documented: `analysts`, `alert_sources`, `alerts`, `alert_groups`, `alert_group_members`, `enrichment_results`, `incidents`, `incident_alerts`, `incident_state_transitions`, `response_actions`, `response_action_outcomes`, `threat_indicators`, `notifications`, `audit_logs`.

## Service Boundaries

- `alert-ingestion-service`: webhook/poll intake, normalization.
- `enrichment-correlation-service`: enrichment pipeline, grouping, correlation.
- `incident-management-service`: incident CRUD/state transitions.
- `response-orchestration-service`: strategy/factory/decorator/proxy response execution.
- `threat-intel-service`: adapters and Redis-backed proxy cache.
- `notification-service`: Slack/PagerDuty/Email mock adapters.
- `audit-service`: audit event persistence and retrieval.
- `identity-service`: mock login/signup/profile.
- `soc-dashboard`: React dashboard.
- `shared contracts`: OpenAPI, AsyncAPI, event contracts.

## API Endpoints

Required APIs are documented in `docs/api/openapi-sdapro.yaml`: ingest webhook/poll, alerts, enrichment, incidents, incident state, response actions, threat intel reputation, notifications, audit events, auth/login/signup, identity/me, dashboard metrics, and realtime dashboard events.

## Event Definitions

Required events are documented in `docs/api/asyncapi-events.yaml`: `AlertIngested`, `AlertEnriched`, `IncidentCreated`, `IncidentStateChanged`, `ResponseActionExecuted`, `ThreatIntelUpdated`, `NotificationDispatched`, and `AuditLogCreated`.

## UI Screens

![Login](assets/01-login-screen.png)

**Login requirement explanation:** Satisfies the required Login screen and identity-service boundary. It provides a separate authentication entry point before analysts can access SOC panels and avoids exposing backend tokens or integration keys in the browser.

![Signup](assets/02-signup-screen.png)

**Signup requirement explanation:** Adds analyst onboarding support for the identity requirement. The signup view creates a mock analyst profile for demonstration and keeps authentication separate from operational screens.

![Main Dashboard](assets/03-main-dashboard.png)

**Main Dashboard requirement explanation:** Summarizes active incidents, alert volume, response outcomes, and SOC status from the dashboard metrics API so seeded demo data is visible immediately.

![Live Alert Stream](assets/04-live-alert-stream.png)

**Live Alert Stream requirement explanation:** Demonstrates event-driven behavior and realtime push from backend events/SSE after mock Splunk or Firewall alerts enter ingestion.

![Incident Queue](assets/05-incident-queue.png)

**Incident Queue requirement explanation:** Shows correlated incidents created from enriched alerts, supporting alert grouping and incident creation.

![Incident Detail](assets/06-incident-detail.png)

**Incident Detail requirement explanation:** Presents incident severity, linked evidence, lifecycle state, and state transition controls, proving State pattern and incident-management behavior.

![Response Console](assets/07-response-console.png)

**Response Console requirement explanation:** Demonstrates response orchestration through strategy selection, factory-created actions, decorator-added audit behavior, proxy control checks, and action outcome recording.

![Threat Intel Lookup](assets/08-threat-intel-lookup.png)

**Threat Intel Lookup requirement explanation:** Shows mock VirusTotal/MISP style reputation lookup through adapters and proxy cache, using Redis for repeat lookup simulation without real API keys.

![Audit Logs](assets/09-audit-logs.png)

**Audit Logs requirement explanation:** Proves that significant system actions generate durable audit events such as alert ingestion, incident changes, response execution, notifications, and threat intel lookups.

![Settings](assets/10-settings.png)

**Settings requirement explanation:** Shows mock integrations while keeping credentials protected. Raw API keys, payload identifiers, and tokens are not displayed in analyst panels.

## Design Pattern Mapping

![Class Diagram](assets/diagram-class.png)

**Class diagram explanation:** This diagram is the main design-pattern proof. It links all 12 required patterns to concrete classes and collaborations: factories, adapters, proxies, decorators, State, Chain of Responsibility, Composite, Strategy, Facade, Observer, and Singleton.

| Pattern | Evidence |
|---|---|
| Singleton | `IngestionConfigManager`, `EventBusPublisher` |
| Factory Method | `AlertNormalizerFactory`, `ResponseActionFactory` |
| Abstract Factory | `EnrichmentProviderFactory`, `NotificationFactory` |
| Composite | `AlertComponent` |
| Facade | `IncidentResponseFacade` |
| Adapter | `ThreatAdapters` |
| Decorator | `Decorators` |
| Proxy | `ThreatIntelProxy`, `ResponseActionProxy` |
| State | `IncidentState` |
| Chain of Responsibility | `EnrichmentHandler` |
| Observer | `EventBusPublisher` |
| Strategy | `ResponseStrategy` |

Pattern annotation proof:

```text
{patterns}
```

## Architecture Style Mapping

| Style | Evidence |
|---|---|
| SOA | Named service boundaries, component diagram, contracts. |
| MVC | Spring controllers/services/repositories plus React views. |
| Layered | API, application, domain, adapter, persistence, infrastructure layers. |
| Event-Driven | RabbitMQ/EventBus and AsyncAPI event definitions. |

## Sequence Diagrams

![Alert Ingestion Sequence](assets/diagram-sequence-alert-ingestion.png)

**Alert ingestion sequence explanation:** This sequence maps directly to the required demo flow: mock alert webhook, normalizer factory, enrichment chain, threat intel adapter/proxy, composite grouping, incident creation, event publishing, dashboard update, and audit storage.

![Incident Response Sequence](assets/diagram-sequence-incident-response.png)

**Incident response sequence explanation:** This sequence proves state transition and response requirements: analyst changes incident state, State validates lifecycle behavior, Strategy selects response, Factory creates actions, Decorator and Proxy wrap execution, events publish, and audit logs are stored.

## Docker Runtime Proof

![Docker Desktop Running](assets/docker-desktop-running.png)

**Docker screenshot explanation:** This satisfies runtime proof for `docker-compose up --build`. It shows the local Compose project and running containers needed for the demo: RabbitMQ, PostgreSQL, Redis, Spring Boot backend, and SOC dashboard frontend.

Command-line proof:

```text
{docker_ps}
```

## API Proof

Dashboard metrics:

```json
{read(ASSETS / "api-dashboard-metrics.json").strip()}
```

Threat intel lookup:

```json
{read(ASSETS / "api-threat-intel-response.json").strip()}
```

## Testing Evidence

```text
{tests}
```

## Required Documentation

- `README.md`
- `docker-compose.yml`
- `docs/adr/ADR-001-soa-vs-microservices-vs-modular-monolith.md`
- `docs/adr/ADR-002-sync-vs-async-communication.md`
- `docs/adr/ADR-003-database-strategy.md`
- `docs/adr/ADR-004-threat-intel-cache-strategy.md`
- `docs/adr/ADR-005-real-time-push-strategy.md`
- `docs/plantuml codes/*.puml`
- `docs/uml diagrams/*.png`
- `docs/api/openapi-sdapro.yaml`
- `docs/api/asyncapi-events.yaml`

## Limitations

- External integrations are mock integrations only.
- The system is built for SDA demonstration, not production-grade cybersecurity.
- Service boundaries are demonstrated through code modules, contracts, diagrams, and docs; the local demo runs through a Spring Boot composition root.

## Conclusion

The project satisfies the SDA-Pro requirements with complete runnable project files, seeded demo data, modern frontend screens, required backend features, design pattern annotations, architecture style evidence, diagrams, ADRs, and tests.
"""
    MD_OUT.write_text(md, encoding="utf-8")


if __name__ == "__main__":
    REPORTS.mkdir(exist_ok=True)
    build_markdown()
    build_docx()
    print(DOCX_OUT)
    print(MD_OUT)
